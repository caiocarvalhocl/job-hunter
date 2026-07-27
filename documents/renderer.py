"""Render documents to DOCX in memory (no disk writes).

Two entry points:
  - text_to_docx: turns a block of text (a cover letter) into a DOCX.
  - cv_to_docx:   turns a *resolved* CV dict (see generators/tailored_cv.py)
                  into an ATS-friendly, single-column DOCX.

Kept intentionally single-column with standard headings: ATS parsers choke on
tables, text boxes and multi-column layouts, so the CV avoids all of them.
"""
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _base_document() -> Document:
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(11)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(48)
        section.left_margin = section.right_margin = Pt(54)
    return doc


def text_to_docx(title: str, body: str, filename: str = "documento.docx") -> BytesIO:
    """Cover-letter style document: a heading plus paragraphs."""
    doc = _base_document()
    doc.add_heading(title, level=1)
    for block in body.split("\n\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block)
    return _to_buffer(doc, filename)


def _section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)


def cv_to_docx(cv: dict, filename: str = "cv.docx") -> BytesIO:
    """Render a resolved CV dict. Expected keys:

    name, contact_lines (list[str]), headings (dict), summary (str),
    skills (list[(label, [items])]), experience (list of dicts with
    header/meta/bullets), projects (list of dicts with header/bullets),
    education (list of (header, meta)), certifications (list[str]),
    languages_label, languages.
    """
    doc = _base_document()

    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(cv["name"])
    name_run.bold = True
    name_run.font.size = Pt(16)

    # Contact
    for line in cv["contact_lines"]:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(line)
        r.font.size = Pt(9.5)

    h = cv["headings"]

    if cv.get("summary"):
        _section_heading(doc, h["summary"])
        doc.add_paragraph(cv["summary"])

    if cv.get("skills"):
        _section_heading(doc, h["skills"])
        for label, items in cv["skills"]:
            p = doc.add_paragraph()
            lab = p.add_run(f"{label}: ")
            lab.bold = True
            p.add_run(", ".join(items))

    if cv.get("experience"):
        _section_heading(doc, h["experience"])
        for exp in cv["experience"]:
            hp = doc.add_paragraph()
            hp.add_run(exp["header"]).bold = True
            if exp.get("meta"):
                mp = doc.add_paragraph()
                mr = mp.add_run(exp["meta"])
                mr.italic = True
                mr.font.size = Pt(9.5)
            for b in exp["bullets"]:
                doc.add_paragraph(b, style="List Bullet")

    if cv.get("projects"):
        _section_heading(doc, h["projects"])
        for proj in cv["projects"]:
            hp = doc.add_paragraph()
            hp.add_run(proj["header"]).bold = True
            for b in proj["bullets"]:
                doc.add_paragraph(b, style="List Bullet")

    if cv.get("education"):
        _section_heading(doc, h["education"])
        for header, meta in cv["education"]:
            hp = doc.add_paragraph()
            hp.add_run(header).bold = True
            if meta:
                mp = doc.add_paragraph()
                mr = mp.add_run(meta)
                mr.italic = True
                mr.font.size = Pt(9.5)

    if cv.get("certifications"):
        _section_heading(doc, h["certifications"])
        for cert in cv["certifications"]:
            doc.add_paragraph(cert, style="List Bullet")

    if cv.get("languages"):
        _section_heading(doc, cv["languages_label"])
        doc.add_paragraph(cv["languages"])

    return _to_buffer(doc, filename)


def _to_buffer(doc: Document, filename: str) -> BytesIO:
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    buf.name = filename
    return buf
